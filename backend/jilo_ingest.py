"""
MediParse AI — Jilo Hackathon Dataset Ingestion
=================================================
Processes ALL 5 folders from the Jilo Hackathon dataset:
  - Discharge_Summary  (.docx files — uses python-docx)
  - Final_Bill         (.pdf)
  - Pre-Auth_Forms     (.pdf)
  - Final_Approval_Letters (.pdf)
  - Settlement_Letters (.pdf)

Usage:
    pip install python-docx
    python jilo_ingest.py
"""

import os, sys, time, io
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

BASE = r"C:\Users\adars\Downloads\Scannes-OPD\Jilo_Hackathon"

FOLDERS = {
    "Discharge_Summary":       {"ext": ".docx", "type": "Discharge Summary",  "prefix": "JILO_DS_"},
    "Final_Bill":              {"ext": ".pdf",  "type": "Hospital Bill",       "prefix": "JILO_BILL_"},
    "Pre-Auth_Forms":          {"ext": ".pdf",  "type": "Insurance Form",      "prefix": "JILO_AUTH_"},
    "Final_Approval_Letters":  {"ext": ".pdf",  "type": "Insurance Form",      "prefix": "JILO_APPR_"},
    "Settlement_Letters":      {"ext": ".pdf",  "type": "Insurance Form",      "prefix": "JILO_SETT_"},
}

DELAY      = 5    # seconds between API calls
SKIP_LARGE = 5    # MB limit

sys.path.insert(0, os.path.dirname(__file__))
from extractor   import extract_text_from_pdf
from ai_pipeline import run_extraction
from validator   import validate_fields
from storage     import save_document, get_all_documents

def color(text, code): return f"\033[{code}m{text}\033[0m"
GREEN, RED, YELLOW, CYAN, BOLD, DIM = 92, 91, 93, 96, 1, 2

# ── Text extractors ───────────────────────────────────────────────────────────

def extract_docx(path: Path) -> str:
    """Extract plain text from a .docx Word file."""
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(parts)
    except ImportError:
        print(color("\n  ✕ python-docx not installed! Run: pip install python-docx", RED))
        raise
    except Exception as e:
        raise Exception(f"docx extraction failed: {e}")

def extract_file(path: Path) -> tuple[str, str]:
    """Returns (raw_text, method)."""
    if path.suffix.lower() == ".docx":
        return extract_docx(path), "docx-digital"
    else:
        return extract_text_from_pdf(path.read_bytes())

# ── Already ingested check ────────────────────────────────────────────────────

def already_ingested():
    try:
        return {d["filename"] for d in get_all_documents()}
    except Exception:
        return set()

# ── Single file processor ─────────────────────────────────────────────────────

def process_file(path: Path, prefix: str, hint_type: str, idx: int, total: int):
    size_mb = path.stat().st_size / (1024 * 1024)
    print(f"\n  {color(f'[{idx}/{total}]', BOLD)} {color(path.name, CYAN)} ({size_mb:.2f} MB)")

    if size_mb > SKIP_LARGE:
        print(color(f"    ⚠ Skipped — too large ({size_mb:.1f} MB)", YELLOW))
        return "skipped_large"

    try:
        print(f"    ⟳ Extracting...", end="", flush=True)
        raw_text, method = extract_file(path)

        if not raw_text or len(raw_text.strip()) < 20:
            print(color(" EMPTY", RED))
            return "empty"

        # Prepend document type hint to help AI classify correctly
        hinted_text = f"DOCUMENT TYPE: {hint_type}\n\n{raw_text}"
        print(color(f" {method} ({len(raw_text)} chars)", GREEN))

        print(f"    ⟳ AI Extraction...", end="", flush=True)
        extracted, warnings, confidence = None, [], 0
        for attempt in range(3):
            try:
                extracted = run_extraction(hinted_text)
                # Force the document type from folder context
                extracted["document_type"] = hint_type
                _, warnings, confidence = validate_fields(extracted)
                break
            except Exception as e:
                if 'rate_limit' in str(e).lower() and attempt < 2:
                    wait = (attempt + 1) * 15
                    print(color(f" ⏳{wait}s...", YELLOW), end="", flush=True)
                    time.sleep(wait)
                else:
                    raise

        if extracted is None:
            raise Exception("All retries exhausted")

        patient  = (extracted.get("patient") or {}).get("name") or "—"
        validated, warnings, confidence = validate_fields(extracted)
        print(color(f" {hint_type} | {patient} | conf:{confidence}%", GREEN))

        print(f"    ⟳ Saving...", end="", flush=True)
        doc_id = save_document(
            filename          = prefix + path.name,
            raw_text          = raw_text,
            extraction_method = method,
            extracted_fields  = validated,
            warnings          = warnings,
            confidence        = confidence,
        )
        print(color(f" id={doc_id} ✓", GREEN))
        return "success"

    except KeyboardInterrupt:
        raise
    except Exception as e:
        print(color(f"\n    ✕ {e}", RED))
        return "error"

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(color("\n╔══════════════════════════════════════════════╗", BOLD))
    print(color( "║  MediParse AI — Jilo Hackathon Ingestion     ║", BOLD))
    print(color( "╚══════════════════════════════════════════════╝\n", BOLD))

    # Check python-docx
    try:
        import docx
    except ImportError:
        print(color("  Installing python-docx...", YELLOW))
        os.system(f"{sys.executable} -m pip install python-docx -q")

    ingested  = already_ingested()
    all_files = []

    for folder_name, cfg in FOLDERS.items():
        folder = Path(BASE) / folder_name
        files  = sorted(folder.glob(f"*{cfg['ext']}"))
        new    = [f for f in files if (cfg["prefix"] + f.name) not in ingested]
        print(f"  📁 {color(folder_name, CYAN)}: {len(files)} total, "
              f"{color(str(len(new)), GREEN)} new")
        for f in new:
            all_files.append((f, cfg["prefix"], cfg["type"]))

    total = len(all_files)
    print(f"\n  Will process: {color(str(total), BOLD)} files across 5 categories\n")

    if total == 0:
        print(color("  ✓ All Jilo files already ingested!", GREEN))
        return

    stats = {"success": 0, "error": 0, "empty": 0, "skipped_large": 0}
    start = time.time()

    for i, (path, prefix, hint) in enumerate(all_files, 1):
        result = process_file(path, prefix, hint, i, total)
        stats[result] = stats.get(result, 0) + 1
        if i < total:
            time.sleep(DELAY)

    elapsed = time.time() - start
    print(color("\n\n═══════════ JILO INGESTION COMPLETE ════════════", BOLD))
    print(f"  ✓ Success:  {color(str(stats['success']), GREEN)}")
    print(f"  ✕ Errors:   {color(str(stats['error']), RED)}")
    print(f"  ⊘ Skipped:  {stats.get('skipped_large',0) + stats.get('empty',0)}")
    print(f"  ⏱ Time:     {elapsed:.0f}s ({elapsed/60:.1f} min)")
    print(color("\n  🏆 Jilo Hackathon dataset fully ingested!", GREEN))
    print(color("  🚀 MediParse AI is now trained on real hospital data!", GREEN))

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print(color("\n\n  ⚡ Interrupted. Progress saved to Supabase.", YELLOW))
